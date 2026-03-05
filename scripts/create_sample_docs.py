"""
scripts/create_sample_docs.py
------------------------------
Generates professional sample documents for all four departments so the
RAG system has rich, realistic content to index on first run.

Documents created
-----------------
HR (hr/)
  ├── vacation_policy.txt
  └── employee_benefits_2026.docx

IT (it/)
  ├── server_security_protocols.pdf
  └── wifi_access_guide.txt

Finance (finance/)
  ├── q1_budget_allocation.xlsx
  └── expense_reimbursement_rules.pdf

general/ (root)
  └── company_history_and_mission.md
"""

import sys
from pathlib import Path

# Resolve project root regardless of where the script is invoked from
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

DOCS_DIR = PROJECT_ROOT / "documents"


# ── Helpers ────────────────────────────────────────────────────────────────────

def ensure_dirs():
    for sub in ("hr", "it", "finance", "general"):
        (DOCS_DIR / sub).mkdir(parents=True, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# HR DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════

def create_vacation_policy_txt():
    content = """\
ACME CORPORATION — EMPLOYEE VACATION POLICY
Policy Number: HR-POL-001 | Effective Date: January 1, 2026 | Version: 3.2
Approved by: VP of Human Resources | Classification: Internal

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. PURPOSE
   This policy establishes a uniform framework for vacation time accrual,
   usage, carryover, and payout for all full-time and part-time employees
   of Acme Corporation.

2. ELIGIBILITY
   All regular full-time employees (≥30 hours/week) are eligible for paid
   vacation benefits beginning on their first day of employment. Part-time
   employees working 20–29 hours per week receive pro-rated benefits.
   Contractors and temporary staff are not covered by this policy.

3. ACCRUAL SCHEDULE
   Vacation time accrues on a bi-weekly basis according to the schedule below:

   Years of Service       Annual Accrual     Bi-Weekly Accrual
   ─────────────────────────────────────────────────────────────
   0 – 2 years            10 days (80 hrs)   3.08 hours
   3 – 5 years            15 days (120 hrs)  4.62 hours
   6 – 10 years           20 days (160 hrs)  6.15 hours
   10+ years              25 days (200 hrs)  7.69 hours
   Senior Directors+      30 days (240 hrs)  9.23 hours

   Accrual begins on the first day of employment but hours are not
   available for use until 90 days of continuous employment have elapsed
   (probationary period).

4. VACATION REQUESTS
   4.1 Employees must submit vacation requests through the HR portal
       (hr.acme.internal) at least:
       – 2 business days in advance for requests ≤3 days
       – 2 weeks in advance for requests of 4–9 days
       – 30 days in advance for requests of 10 days or more
   4.2 All requests are subject to manager approval. Approval is contingent
       on business needs and adequate team coverage.
   4.3 Managers must respond to requests within 5 business days. No response
       within that window constitutes automatic approval for requests ≤5 days.

5. CARRYOVER RULES
   5.1 Employees may carry over a maximum of 5 days (40 hours) of unused
       vacation into the following calendar year.
   5.2 Carryover must be used by March 31 of the new year or it will be
       forfeited. There is no cash payout for forfeited carryover.
   5.3 Exceptions to the 5-day cap may be granted by the HR Director for
       employees on approved medical, parental, or military leave.

6. VACATION PAYOUT
   6.1 Upon voluntary resignation with ≥2 weeks' notice, employees will
       receive a lump-sum payout for all accrued, unused vacation.
   6.2 Upon involuntary termination, accrued vacation is paid per applicable
       state law. In states where payout is mandated (CA, MA, ND), full
       accrued balances are paid regardless of tenure.
   6.3 Employees terminated for gross misconduct forfeit accrued vacation
       except where prohibited by law.

7. BLACKOUT PERIODS
   Vacation requests during the following windows require VP-level approval:
   – December 23 – January 2 (Year-End Close)
   – Last 3 business days of each fiscal quarter
   – Any period declared a "Critical Business Period" by Executive Leadership

8. SPECIAL PROVISIONS
   8.1 Bereavement: Employees are entitled to up to 5 paid days for the
       death of an immediate family member (spouse, child, parent, sibling).
       This is separate from vacation and does not count against the vacation
       balance.
   8.2 Jury Duty: Paid leave for jury duty up to 30 days. Does not count
       against vacation balance.
   8.3 Floating Holidays: Each employee receives 2 floating holidays per year
       in addition to vacation. These do not accrue and must be used within
       the calendar year.

9. POLICY VIOLATIONS
   Misrepresentation of vacation usage, unapproved absences charged to
   vacation, or falsification of time records may result in disciplinary
   action up to and including termination.

10. POLICY REVIEW
    This policy is reviewed annually by the HR Department. Employees will
    be notified of material changes via company email and the HR portal.

For questions, contact HR at hr-support@acme.internal or ext. 4400.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
© 2026 Acme Corporation. All rights reserved. Internal Use Only.
"""
    path = DOCS_DIR / "hr" / "vacation_policy.txt"
    path.write_text(content, encoding="utf-8")
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


def create_employee_benefits_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Acme Corporation — Employee Benefits Guide 2026", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Effective: January 1, 2026  |  Document ID: HR-BEN-2026  |  Classification: Internal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Introduction
    doc.add_heading("1. Introduction", 1)
    doc.add_paragraph(
        "Acme Corporation is committed to the health, financial security, and overall "
        "well-being of its employees. This guide outlines all benefits available to eligible "
        "full-time employees for the 2026 benefit year. Open enrollment runs November 1–30, 2025. "
        "Changes take effect January 1, 2026. For assistance, contact benefits@acme.internal."
    )

    # Health Insurance
    doc.add_heading("2. Medical Insurance", 1)
    doc.add_paragraph(
        "Acme offers three medical plan tiers through BlueCross BlueShield of America. "
        "All plans cover preventive care at 100% with in-network providers."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Plan", "Monthly Premium (EE)", "Monthly Premium (Family)", "Deductible", "Out-of-Pocket Max"]):
        hdr[i].text = h
    data = [
        ("Gold Plus",   "$180",  "$480",  "$500",   "$3,000"),
        ("Silver Care", "$110",  "$295",  "$1,500", "$6,000"),
        ("Bronze Basic","$55",   "$145",  "$3,500", "$8,700"),
    ]
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

    # Dental
    doc.add_heading("3. Dental Insurance", 1)
    doc.add_paragraph(
        "Dental coverage is provided through Delta Dental PPO. Preventive services (cleanings, "
        "X-rays) are covered at 100%. Basic restorative services are covered at 80% after "
        "deductible. Major services (crowns, bridges) at 50% after deductible. "
        "Annual maximum benefit: $2,000 per person. Orthodontia: $1,500 lifetime maximum. "
        "Employee monthly premium: $12. Family monthly premium: $38."
    )

    # Vision
    doc.add_heading("4. Vision Insurance", 1)
    doc.add_paragraph(
        "Vision coverage through VSP Choice. Annual eye exam: $0 copay. "
        "Frames: $150 allowance every 24 months. Contact lenses: $150 allowance in lieu of frames. "
        "Employee monthly premium: $6. Family monthly premium: $14."
    )

    # 401(k)
    doc.add_heading("5. 401(k) Retirement Plan", 1)
    doc.add_paragraph(
        "Acme Corporation offers a 401(k) plan administered by Fidelity Investments. "
        "Employees may contribute up to the IRS annual limit ($23,500 for 2026; $31,000 "
        "for employees aged 50+). Acme matches 100% of the first 4% of eligible compensation "
        "contributed, plus 50% of the next 2%. Full vesting schedule:"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Year 1: 0% vested")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Year 2: 25% vested")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Year 3: 50% vested")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Year 4: 75% vested")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Year 5+: 100% vested")

    # Life Insurance
    doc.add_heading("6. Life and Disability Insurance", 1)
    doc.add_paragraph(
        "Basic Life Insurance: 2× annual salary, company-paid. "
        "Supplemental Life: Up to 5× salary (employee-paid, evidence of insurability required for >3×). "
        "Short-Term Disability: 60% of salary for up to 12 weeks after a 7-day waiting period. "
        "Long-Term Disability: 60% of salary after 90-day elimination period, up to age 65."
    )

    # Wellness
    doc.add_heading("7. Wellness Benefits", 1)
    doc.add_paragraph(
        "Gym Reimbursement: Up to $60/month reimbursed for gym membership or fitness classes. "
        "Submit receipts via the Concur expense system under category 'Wellness Reimbursement'. "
        "Mental Health: 12 free counseling sessions per year through our EAP (Employee Assistance "
        "Program) with Lyra Health. Completely confidential. Call 1-800-555-EAP1."
    )

    # PTO Summary
    doc.add_heading("8. Paid Time Off Summary", 1)
    doc.add_paragraph(
        "In addition to vacation (see HR-POL-001), employees receive: 11 company holidays, "
        "2 floating holidays, 10 days sick leave per year (non-accruing), and 5 days "
        "bereavement leave for immediate family."
    )

    # Contact
    doc.add_heading("9. Contact & Resources", 1)
    doc.add_paragraph(
        "Benefits Portal: benefits.acme.internal\n"
        "Benefits Team Email: benefits@acme.internal\n"
        "Benefits Hotline: ext. 4455 (Mon–Fri, 9 AM – 5 PM EST)\n"
        "Fidelity 401(k): 1-800-343-3548\n"
        "EAP (Lyra Health): 1-800-555-EAP1 (24/7)"
    )

    path = DOCS_DIR / "hr" / "employee_benefits_2026.docx"
    doc.save(str(path))
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


# ══════════════════════════════════════════════════════════════════════════════
# IT DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════

def create_server_security_pdf():
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    path = DOCS_DIR / "it" / "server_security_protocols.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=colors.HexColor("#1e3a5f"), spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#2563eb"), spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["Normal"], spaceAfter=8, leading=16)
    bullet = ParagraphStyle("Bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=5, leading=14)

    story = []

    # Header
    story.append(Paragraph("ACME CORPORATION — SERVER SECURITY PROTOCOLS", h1))
    story.append(Paragraph(
        "Document: IT-SEC-2026-001 | Classification: Confidential | Effective: Jan 1, 2026", body
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#2563eb")))
    story.append(Spacer(1, 12))

    # Section 1
    story.append(Paragraph("1. SCOPE AND PURPOSE", h2))
    story.append(Paragraph(
        "This document defines the mandatory security protocols for all servers, virtual machines, "
        "containers, and cloud workloads operated by Acme Corporation's IT Infrastructure team. "
        "Compliance is mandatory for all IT staff. Violations must be reported to the CISO within 24 hours.", body
    ))

    # Section 2 — Access Control
    story.append(Paragraph("2. ACCESS CONTROL", h2))
    story.append(Paragraph("2.1 Privileged Access Management (PAM)", h2))
    for item in [
        "All root/administrator access must be performed via the CyberArk PAM solution.",
        "Shared service accounts are prohibited. Each administrator must use a personal privileged account.",
        "SSH keys must be 4096-bit RSA or Ed25519. Password-based SSH authentication is disabled on all servers.",
        "MFA (TOTP or hardware token) is required for all privileged sessions.",
        "Privileged sessions are recorded and retained for 90 days in the SIEM.",
    ]:
        story.append(Paragraph(f"• {item}", bullet))

    story.append(Paragraph("2.2 Least Privilege Principle", h2))
    story.append(Paragraph(
        "All service accounts must be granted only the minimum permissions necessary. "
        "Access reviews are conducted quarterly. Unused accounts are disabled after 60 days of inactivity.", body
    ))

    # Section 3 — Patch Management
    story.append(Paragraph("3. PATCH MANAGEMENT", h2))
    data = [
        ["Severity",      "Patch Window",   "Max Time to Patch", "Exceptions"],
        ["Critical (9-10)","Emergency (<72h)","72 hours",          "CISO approval required"],
        ["High (7-8.9)",  "Next patch cycle","14 days",           "Manager + CISO approval"],
        ["Medium (4-6.9)","Monthly cycle",   "30 days",           "Manager approval"],
        ["Low (0-3.9)",   "Quarterly cycle", "90 days",           "Team lead approval"],
    ]
    tbl = Table(data, colWidths=[1.3*inch, 1.4*inch, 1.4*inch, 2.2*inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4ff")]),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))

    # Section 4 — Firewall
    story.append(Paragraph("4. FIREWALL AND NETWORK SEGMENTATION", h2))
    for item in [
        "All production servers reside in the PROD VLAN (10.10.0.0/16). No direct internet access.",
        "DMZ servers (10.20.0.0/24) handle public-facing traffic only. They cannot initiate connections to PROD.",
        "Default firewall policy is DENY ALL. Exceptions must be approved via the firewall change request (FCR) process.",
        "All firewall rules must include a business justification and expiry date (max 1 year).",
        "IDS/IPS (Palo Alto PA-5260) monitors all east-west traffic. Alerts are routed to the SOC 24/7.",
    ]:
        story.append(Paragraph(f"• {item}", bullet))

    # Section 5 — Encryption
    story.append(Paragraph("5. ENCRYPTION STANDARDS", h2))
    story.append(Paragraph(
        "Data at Rest: AES-256-GCM for all databases and file storage. "
        "Encryption keys are managed in HashiCorp Vault. Key rotation occurs every 365 days or on demand after a breach. "
        "Data in Transit: TLS 1.3 minimum. TLS 1.0 and 1.1 are disabled on all endpoints. "
        "Certificates are issued by the internal PKI (ACME-CA-ROOT) with a 398-day maximum validity.", body
    ))

    # Section 6 — Incident Response
    story.append(Paragraph("6. INCIDENT RESPONSE", h2))
    story.append(Paragraph(
        "Security incidents must be reported to soc@acme.internal or the 24/7 hotline: ext. 9999. "
        "P1 incidents (data breach, ransomware, APT) require notification to the CISO within 1 hour "
        "and initiation of the Incident Response Playbook (IR-PLAY-001). "
        "All incidents are tracked in ServiceNow under the Security Incident category.", body
    ))

    # Section 7 — Audit
    story.append(Paragraph("7. LOGGING AND AUDIT", h2))
    for item in [
        "All servers must forward logs to the central SIEM (Splunk Enterprise, siem.acme.internal).",
        "Minimum log retention: 1 year online, 7 years in cold storage (AWS S3 Glacier).",
        "Authentication events, privilege escalations, and configuration changes are mandatory log sources.",
        "Log integrity is protected via WORM storage and SHA-256 checksums.",
    ]:
        story.append(Paragraph(f"• {item}", bullet))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "Approved by: Alice Johnson, IT Director | Last Review: December 2025 | "
        "Next Review: December 2026 | Questions: security@acme.internal", body
    ))

    doc.build(story)
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


def create_wifi_access_guide_txt():
    content = """\
ACME CORPORATION — WI-FI ACCESS GUIDE
Document: IT-NET-WIFI-001 | Version: 2.1 | Updated: January 2026
Classification: Public (may be shared with visitors)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

AVAILABLE NETWORKS AT ACME CORPORATION HQ

1. ACME-CORP (Employee Network)
   ─────────────────────────────
   SSID     : ACME-CORP
   Security : WPA3-Enterprise (802.1X)
   Access   : Full corporate intranet, internet, and cloud services
   How to connect:
     a) Select "ACME-CORP" from your device's Wi-Fi list.
     b) When prompted, enter your corporate credentials:
        Username : your_employee_id@acme.internal  (e.g. EMP001@acme.internal)
        Password : Your Windows/Active Directory password
     c) Accept the ACME-CA certificate when prompted (first-time only).
   Speed    : Up to 1 Gbps (Wi-Fi 6 / 802.11ax)
   Coverage : All floors, conference rooms, and parking garage Level B1

2. ACME-GUEST (Visitor Network)
   ─────────────────────────────
   SSID     : ACME-GUEST
   Security : WPA2-PSK
   Password : AcmeVisitor2026!   (rotated every 90 days; ask reception for current password)
   Access   : Internet only. No access to corporate intranet or internal systems.
   Speed    : Up to 50 Mbps per device (bandwidth-limited)
   Duration : Sessions expire after 8 hours. Reconnect with the same password.

3. ACME-IOT (IoT and Conference Room Devices)
   ──────────────────────────────────────────
   SSID     : ACME-IOT
   Access   : Restricted. Used for printers, AV equipment, and smart devices.
   Note     : Employees should NOT connect personal or company laptops to this network.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CONNECTING PERSONAL DEVICES
Personal phones and tablets may be connected to ACME-CORP using your corporate
credentials. However, connecting to ACME-CORP grants IT the right to remotely
wipe the device if it is lost or compromised (per the BYOD Policy, IT-POL-BYOD-003).

For devices you do not want enrolled in MDM, please use ACME-GUEST instead.

TROUBLESHOOTING
  Problem: Cannot authenticate to ACME-CORP
  Solution: Ensure your AD password has not expired. Reset at: https://pwd.acme.internal
            or contact the helpdesk (ext. 4000 or helpdesk@acme.internal).

  Problem: Slow connection speed
  Solution: Avoid ACME-GUEST for large file transfers. Use wired Ethernet
            (available at most desks — Cat6a ports labeled "DATA").

  Problem: Certificate warning on macOS/iOS
  Solution: Install the ACME-CA certificate from https://it.acme.internal/ca-cert.crt

REMOTE WORK — VPN ACCESS
When working outside the office, use the GlobalProtect VPN client to access
internal resources. Download from: https://vpn.acme.internal
Credentials: same as your AD login. MFA via Duo Security is required.
VPN Server: vpn.acme.internal | Split tunneling: disabled for security.

DATA USAGE POLICY
Use of the corporate Wi-Fi for personal streaming, torrenting, or any activity
that violates the Acceptable Use Policy (IT-POL-AUP-001) is prohibited and
may result in disciplinary action.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
IT Helpdesk: ext. 4000 | helpdesk@acme.internal | 8 AM – 8 PM EST, Mon–Fri
"""
    path = DOCS_DIR / "it" / "wifi_access_guide.txt"
    path.write_text(content, encoding="utf-8")
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


# ══════════════════════════════════════════════════════════════════════════════
# FINANCE DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════

def create_q1_budget_xlsx():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference

    wb = openpyxl.Workbook()

    # ── Sheet 1: Budget Summary ────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Q1 Budget Summary"

    header_fill = PatternFill("solid", fgColor="1e3a5f")
    subhdr_fill = PatternFill("solid", fgColor="2563eb")
    alt_fill    = PatternFill("solid", fgColor="dbeafe")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    title_font  = Font(bold=True, size=14, color="1e3a5f")
    bold_font   = Font(bold=True)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    money_fmt = '"$"#,##0.00'

    ws["A1"] = "ACME CORPORATION — Q1 2026 BUDGET ALLOCATION"
    ws["A1"].font = title_font
    ws["A2"] = "Fiscal Quarter: January 1 – March 31, 2026  |  Approved: December 15, 2025  |  Classification: Confidential"
    ws["A2"].font = Font(italic=True, size=10, color="555555")
    ws.merge_cells("A1:G1")
    ws.merge_cells("A2:G2")

    headers = ["Department", "Budget Category", "Q1 Allocated ($)", "Q1 Spent ($)",
               "Q1 Remaining ($)", "% Used", "Status"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = thin

    budget_data = [
        # Dept,         Category,               Allocated,    Spent
        ("HR",          "Salaries & Wages",      1_200_000,   1_195_000),
        ("HR",          "Recruitment",              85_000,      62_400),
        ("HR",          "Training & Development",   45_000,      38_200),
        ("HR",          "Employee Benefits Admin",  22_000,      21_800),
        ("IT",          "Salaries & Wages",        980_000,     978_500),
        ("IT",          "Cloud Infrastructure",    150_000,     132_400),
        ("IT",          "Software Licenses",        95_000,      89_300),
        ("IT",          "Hardware Procurement",     60_000,      48_750),
        ("IT",          "Security Tools",           40_000,      39_800),
        ("Finance",     "Salaries & Wages",        620_000,     619_200),
        ("Finance",     "Audit & Compliance",       55_000,      48_000),
        ("Finance",     "Financial Systems",        30_000,      28_500),
        ("Marketing",   "Salaries & Wages",        540_000,     539_100),
        ("Marketing",   "Digital Advertising",     200_000,     187_600),
        ("Marketing",   "Events & Conferences",     75_000,      42_000),
        ("Operations",  "Salaries & Wages",        890_000,     887_400),
        ("Operations",  "Facilities",              120_000,     114_300),
        ("Operations",  "Equipment Maintenance",    35_000,      28_900),
        ("R&D",         "Salaries & Wages",        750_000,     748_000),
        ("R&D",         "Research Materials",       80_000,      51_200),
        ("R&D",         "External Partnerships",    60_000,      35_000),
    ]

    for row_idx, (dept, cat, allocated, spent) in enumerate(budget_data, 5):
        remaining = allocated - spent
        pct = (spent / allocated) * 100 if allocated else 0
        status = "On Track" if pct <= 95 else ("Over Budget" if pct > 100 else "At Risk")
        row_fill = alt_fill if (row_idx % 2 == 0) else None

        row_vals = [dept, cat, allocated, spent, remaining, pct / 100, status]
        for col, val in enumerate(row_vals, 1):
            cell = ws.cell(row=row_idx, column=col, value=val)
            cell.border = thin
            if row_fill:
                cell.fill = row_fill
            if col in (3, 4, 5):
                cell.number_format = money_fmt
            if col == 6:
                cell.number_format = "0.0%"
            if col == 7:
                if val == "Over Budget":
                    cell.font = Font(bold=True, color="CC0000")
                elif val == "At Risk":
                    cell.font = Font(bold=True, color="CC6600")
                else:
                    cell.font = Font(bold=True, color="006600")

    # Totals row
    total_row = len(budget_data) + 5
    ws.cell(total_row, 1, "TOTAL").font = bold_font
    ws.cell(total_row, 2, "All Departments").font = bold_font
    for col, formula in [(3, f"=SUM(C5:C{total_row-1})"),
                          (4, f"=SUM(D5:D{total_row-1})"),
                          (5, f"=SUM(E5:E{total_row-1})")]:
        cell = ws.cell(total_row, col, formula)
        cell.font = bold_font
        cell.fill = subhdr_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.number_format = money_fmt
        cell.border = thin

    # Column widths
    for col, w in enumerate([18, 28, 18, 16, 18, 10, 14], 1):
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A5"

    # ── Sheet 2: Dept Breakdown ────────────────────────────────────────────────
    ws2 = wb.create_sheet("Department Breakdown")
    ws2["A1"] = "Department Budget Overview — Q1 2026"
    ws2["A1"].font = title_font

    headers2 = ["Department", "Total Allocated ($)", "Total Spent ($)", "Remaining ($)", "Utilization %"]
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin

    dept_totals: dict = {}
    for dept, _, allocated, spent in budget_data:
        if dept not in dept_totals:
            dept_totals[dept] = [0, 0]
        dept_totals[dept][0] += allocated
        dept_totals[dept][1] += spent

    for r, (dept, (alloc, spent)) in enumerate(dept_totals.items(), 4):
        rem = alloc - spent
        pct = spent / alloc if alloc else 0
        for col, val in enumerate([dept, alloc, spent, rem, pct], 1):
            cell = ws2.cell(r, col, val)
            cell.border = thin
            if col in (2, 3, 4):
                cell.number_format = money_fmt
            if col == 5:
                cell.number_format = "0.0%"

    for col, w in enumerate([18, 22, 18, 18, 16], 1):
        ws2.column_dimensions[get_column_letter(col)].width = w

    # ── Sheet 3: Notes ─────────────────────────────────────────────────────────
    ws3 = wb.create_sheet("Budget Notes")
    notes = [
        ("CFO Notes — Q1 2026", True),
        ("", False),
        ("1. Overall Q1 spending is within 3% of budget across all departments.", False),
        ("2. IT Cloud Infrastructure spend of $132,400 against a $150,000 budget reflects the delayed", False),
        ("   migration of the legacy ERP system, now rescheduled to Q2.", False),
        ("3. Marketing Events budget is 56% utilized due to the postponement of the Q1 Industry Summit.", False),
        ("   The remaining $33,000 will be reallocated to Q2 digital campaigns pending CFO approval.", False),
        ("4. R&D External Partnerships shows 58.3% utilization; one partnership agreement ($25,000)", False),
        ("   is under legal review and will be finalized in April.", False),
        ("5. All departments must submit Q2 forecasts by March 15, 2026 to finance@acme.internal.", False),
        ("", False),
        ("Budget Owners:", True),
        ("HR Budget Owner      : Bob Martinez (bob.martinez@acme.internal)", False),
        ("IT Budget Owner      : Alice Johnson (alice.johnson@acme.internal)", False),
        ("Finance Budget Owner : Carol White (carol.white@acme.internal)", False),
        ("CFO                  : James Harrington (j.harrington@acme.internal)", False),
    ]
    for row_n, (txt, bold) in enumerate(notes, 1):
        cell = ws3.cell(row_n, 1, txt)
        if bold:
            cell.font = Font(bold=True, size=12 if row_n == 1 else 11)
    ws3.column_dimensions["A"].width = 80

    path = DOCS_DIR / "finance" / "q1_budget_allocation.xlsx"
    wb.save(str(path))
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


def create_expense_reimbursement_pdf():
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

    path = DOCS_DIR / "finance" / "expense_reimbursement_rules.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=colors.HexColor("#1e3a5f"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#2563eb"))
    body = ParagraphStyle("Body", parent=styles["Normal"], spaceAfter=8, leading=16)
    bullet = ParagraphStyle("Bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=5, leading=14)

    story = []
    story.append(Paragraph("ACME CORPORATION — EXPENSE REIMBURSEMENT POLICY", h1))
    story.append(Paragraph(
        "Policy: FIN-EXP-2026 | Effective: January 1, 2026 | Classification: Internal", body
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#2563eb")))
    story.append(Spacer(1, 10))

    story.append(Paragraph("1. PURPOSE", h2))
    story.append(Paragraph(
        "This policy governs the reimbursement of business expenses incurred by Acme Corporation "
        "employees in the performance of their duties. All expenses must be legitimate, necessary, "
        "reasonable, and properly documented.", body
    ))

    story.append(Paragraph("2. APPROVAL LIMITS", h2))
    limits_data = [
        ["Employee Level",         "No Approval Needed", "Manager Approval", "VP Approval", "CFO Approval"],
        ["Individual Contributor", "Up to $50",          "$51 – $500",       "$501 – $2,000","$2,001+"],
        ["Team Lead / Analyst",    "Up to $100",         "$101 – $1,000",    "$1,001–$5,000","$5,001+"],
        ["Manager / Director",     "Up to $500",         "$501 – $2,500",    "$2,501–$10,000","$10,001+"],
        ["VP / SVP",               "Up to $2,000",       "$2,001 – $10,000", "$10,001–$25,000","$25,001+"],
    ]
    tbl = Table(limits_data, colWidths=[1.6*inch, 1.1*inch, 1.1*inch, 1.3*inch, 1.1*inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4ff")]),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))

    story.append(Paragraph("3. ELIGIBLE EXPENSE CATEGORIES", h2))
    cats = [
        ("Business Travel", "$0.70/mile (personal vehicle); actual cost for flights booked ≥14 days in advance via Concur. Business class only for flights >8 hours."),
        ("Meals (Solo Business)", "Up to $75/day when traveling. No alcohol reimbursement."),
        ("Client Entertainment", "Up to $150/person with client present. Itemized receipt required. Manager pre-approval for >$500 total."),
        ("Hotel Accommodation", "Up to $250/night in standard cities; $350/night in high-cost cities (NYC, SF, London, Tokyo). See city rate schedule FIN-CITY-RATES."),
        ("Software/Subscriptions", "Requires IT pre-approval. See IT-POL-SAAS-001 for approved vendor list."),
        ("Office Supplies", "Up to $100/month without approval. Submit receipts via Concur."),
        ("Professional Development", "Conferences, courses, and books up to $2,000/year per employee. Manager approval required."),
        ("Home Office (Remote Employees)", "One-time setup allowance of $500. Monthly internet stipend: $50."),
    ]
    for cat, detail in cats:
        story.append(Paragraph(f"<b>{cat}:</b> {detail}", bullet))

    story.append(Paragraph("4. INELIGIBLE EXPENSES", h2))
    ineligible = [
        "Personal grooming, clothing, or accessories",
        "Traffic/parking fines and violations",
        "Personal entertainment or leisure activities",
        "Gifts for personal relationships (birthday gifts, wedding gifts, etc.)",
        "Upgrades to first class without CFO approval",
        "Spouse/partner travel unless pre-approved as a business requirement",
        "Expenses incurred more than 90 days prior to submission",
        "Alcohol at client events unless manager-approved with signed justification",
    ]
    for item in ineligible:
        story.append(Paragraph(f"• {item}", bullet))

    story.append(Paragraph("5. SUBMISSION PROCESS", h2))
    steps = [
        "Log into Concur Expense at concur.acme.internal using your SSO credentials.",
        "Create a new Expense Report. Attach receipts (photos accepted via the Concur mobile app).",
        "Select the correct expense category and project code (obtain from your manager).",
        "Submit within 30 days of the expense date. Late submissions (31–90 days) require manager justification.",
        "Reimbursements are processed within 10 business days and paid via ACH to your registered bank account.",
    ]
    for i, step in enumerate(steps, 1):
        story.append(Paragraph(f"{i}. {step}", bullet))

    story.append(Paragraph("6. RECEIPT REQUIREMENTS", h2))
    story.append(Paragraph(
        "Receipts are required for ALL expenses over $25. Corporate credit card statements alone "
        "are NOT acceptable substitutes for itemized receipts. Receipts must show: vendor name, "
        "date, itemized amounts, and payment method. For meals, the receipt must include the "
        "list of attendees and business purpose.", body
    ))

    story.append(Paragraph("7. NON-COMPLIANCE", h2))
    story.append(Paragraph(
        "Fraudulent expense claims will result in disciplinary action up to and including "
        "termination and possible criminal prosecution. All expense reports are subject to "
        "random audit by the Internal Audit team. Employees with >3 policy violations in 12 months "
        "will be required to pre-approve all expenses with the Finance Controller.", body
    ))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "Questions? Contact: expenses@acme.internal | ext. 4200 | "
        "Finance team office hours: Mon–Fri 9 AM – 4 PM EST", body
    ))

    doc.build(story)
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


# ══════════════════════════════════════════════════════════════════════════════
# GENERAL DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def create_company_history_md():
    content = """\
# Acme Corporation — Company History and Mission

**Document Type:** Public Reference | **Last Updated:** January 2026 | **Owner:** Communications Team

---

## Our Story

Acme Corporation was founded in **1987** by Dr. Eleanor Hayes and Marcus Trent in a small office in
Austin, Texas. What began as a two-person engineering consulting firm has grown into a global
technology and services enterprise with over **14,000 employees** across **32 countries**.

### 1987–1995: The Founding Years

Dr. Hayes, a computer scientist from MIT, and Marcus Trent, a mechanical engineer from Stanford,
met at a technology conference in 1985. Bonded by a shared vision of making enterprise software
more accessible to mid-sized businesses, they pooled their savings to incorporate Acme Corporation
on **March 14, 1987**. Their first product, *AcmeLink*, was a pioneering data exchange platform
that enabled companies to integrate their siloed databases — groundbreaking technology for the era.

By 1992, AcmeLink had been adopted by over 400 companies across the United States. The company
reached its first $10 million in revenue in 1993 and completed its Series A funding round in 1994,
raising $8.5 million led by Sequoia Capital.

### 1996–2005: Global Expansion

The late 1990s saw Acme expand aggressively into Europe and Asia-Pacific. The company opened its
first international office in **London in 1997**, followed by Singapore (1999), Tokyo (2001),
and Sydney (2003).

In **2000**, Acme went public on the NASDAQ (ticker: ACME), raising $240 million in its IPO.
The funds were used to acquire three strategic companies:
- **DataBridge Systems** (2001) — enterprise middleware
- **Fortis Analytics** (2003) — business intelligence platform
- **NovaSec Ltd.** (2004) — cybersecurity solutions

### 2006–2015: The Cloud Transition

Recognizing the shift to cloud computing, Acme launched its **Acme Cloud Platform (ACP)** in 2008.
Within three years, ACP had become the company's fastest-growing product line. The company
surpassed **$1 billion in annual revenue** for the first time in **2011**.

During this period, Acme also deepened its commitment to corporate responsibility, launching:
- The **Acme Foundation** (2009): $5M/year in STEM education grants
- The **GreenOps Initiative** (2012): commitment to carbon-neutral operations by 2030

### 2016–Present: AI and Innovation

Under the leadership of CEO **Sandra Reeves** (appointed 2016), Acme has positioned itself at the
forefront of enterprise AI. Key milestones include:

- **2018**: Launch of AcmeAI, an ML-powered analytics suite
- **2020**: Acquisition of Luminary.ai for $310 million
- **2022**: Acme named to Fortune 500 (#312)
- **2024**: Revenue exceeds **$4.2 billion**; employee headcount reaches 14,000+
- **2025**: Launch of **Acme Enterprise Knowledge Platform**, integrating AI across all product lines

---

## Our Mission

> *"To empower organizations with intelligent, secure, and scalable technology that transforms data
> into decisions and decisions into outcomes."*

We pursue this mission through three strategic pillars:

### 1. Innovation Without Compromise
We invest **12% of annual revenue** in R&D — above the industry average — because we believe
tomorrow's problems require today's investment in foundational research.

### 2. Security by Design
Every product we build is designed with security as a first principle, not an afterthought.
Our customers trust us with their most sensitive data, and we take that responsibility seriously.

### 3. Human-Centered Technology
Technology should serve people, not the other way around. We design with empathy, build for
accessibility, and measure success by the outcomes our customers achieve.

---

## Core Values

| Value            | What It Means at Acme                                                 |
|------------------|-----------------------------------------------------------------------|
| **Integrity**    | We do what we say. We say what we mean. No exceptions.                |
| **Curiosity**    | We ask "why?" and "what if?" before "how do we maintain the status quo?" |
| **Accountability** | We own our results — the wins and the misses.                       |
| **Inclusion**    | Every voice belongs. Diversity of thought drives better outcomes.     |
| **Excellence**   | Good enough never is. We pursue mastery in everything we do.         |

---

## Leadership Team (as of January 2026)

- **Sandra Reeves** — Chief Executive Officer
- **James Harrington** — Chief Financial Officer
- **Dr. Priya Nair** — Chief Technology Officer
- **Marcus Trent** — Chairman of the Board (Co-Founder)
- **Linda Osei** — Chief People Officer
- **Rafael Santos** — Chief Revenue Officer
- **Alice Johnson** — VP of IT Infrastructure

---

## Offices and Locations

**Global Headquarters:** 1 Acme Plaza, Austin, TX 78701, USA
**Regional HQs:** London | Singapore | Tokyo | São Paulo | Dubai
**R&D Centers:** Austin | Boston | Berlin | Bangalore

---

## Key Statistics (FY 2025)

| Metric                     | Value            |
|----------------------------|------------------|
| Annual Revenue             | $4.2 Billion     |
| Employees Worldwide        | 14,200           |
| Countries of Operation     | 32               |
| Active Enterprise Clients  | 8,500+           |
| Patents Held               | 340+             |
| R&D Investment (% Revenue) | 12%              |
| Carbon Offset Achieved     | 78% (target: 100% by 2030) |

---

*For media inquiries: press@acme.com | Investor Relations: ir@acme.com*
*© 2026 Acme Corporation. All rights reserved.*
"""
    path = DOCS_DIR / "general" / "company_history_and_mission.md"
    path.write_text(content, encoding="utf-8")
    print(f"  ✅ Created: {path.relative_to(PROJECT_ROOT)}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("\n🏗️  Generating Enterprise Knowledge Brain sample documents…\n")
    ensure_dirs()

    print("📁 HR Department:")
    create_vacation_policy_txt()
    create_employee_benefits_docx()

    print("\n📁 IT Department:")
    create_server_security_pdf()
    create_wifi_access_guide_txt()

    print("\n📁 Finance Department:")
    create_q1_budget_xlsx()
    create_expense_reimbursement_pdf()

    print("\n📁 General:")
    create_company_history_md()

    print(f"\n✨ All 7 documents created under: {DOCS_DIR.relative_to(PROJECT_ROOT)}/")
    print("""
Directory structure:
  documents/
  ├── hr/
  │   ├── vacation_policy.txt
  │   └── employee_benefits_2026.docx
  ├── it/
  │   ├── server_security_protocols.pdf
  │   └── wifi_access_guide.txt
  ├── finance/
  │   ├── q1_budget_allocation.xlsx
  │   └── expense_reimbursement_rules.pdf
  └── general/
      └── company_history_and_mission.md

Next steps:
  1. pip install -r requirements.txt
  2. ollama pull llama3.2
  3. python scripts/create_sample_docs.py   (already done!)
  4. streamlit run src/app.py
     → In the Admin panel, click "Bulk Ingest Sample Docs"
     → Select an employee and start chatting!
""")


if __name__ == "__main__":
    main()
